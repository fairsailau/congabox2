"""
Utility for parsing DOCX files to extract text and Conga merge fields.
"""
import re
import docx
from typing import Dict, List, Tuple

class DocxParser:
    """Parser for extracting content from Conga template DOCX files."""
    
    def __init__(self):
        """Initialize the DOCX parser."""
        # Regex pattern for Conga merge fields
        self.conga_field_pattern = re.compile(r'«([^»]+)»')
    
    def parse_template(self, docx_file) -> Tuple[str, List[Dict]]:
        """
        Parse a Conga template DOCX file to extract text and merge fields.
        
        Args:
            docx_file: Path to the DOCX file or file-like object
            
        Returns:
            Tuple containing:
                - Full text content of the document
                - List of dictionaries with field information
        """
        try:
            doc = docx.Document(docx_file)
            
            # Extract full text
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_text += para.text + "\n"
            
            # Extract merge fields
            merge_fields = []
            field_matches = self.conga_field_pattern.findall(full_text)
            
            for field in field_matches:
                field_info = {
                    "original": f"«{field}»",
                    "name": field,
                    "context": self._get_field_context(full_text, field)
                }
                merge_fields.append(field_info)
            
            return full_text, merge_fields
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")
    
    def _get_field_context(self, text: str, field: str, context_chars: int = 50) -> str:
        """
        Get surrounding context for a merge field.
        
        Args:
            text: Full document text
            field: Field name without delimiters
            context_chars: Number of characters to include before and after
            
        Returns:
            String with context before and after the field
        """
        field_with_delimiters = f"«{field}»"
        field_pos = text.find(field_with_delimiters)
        
        if field_pos == -1:
            return ""
        
        start = max(0, field_pos - context_chars)
        end = min(len(text), field_pos + len(field_with_delimiters) + context_chars)
        
        context = text[start:end]
        if start > 0:
            context = "..." + context
        if end < len(text):
            context = context + "..."
            
        return context
