"""
Sample Conga template for testing purposes.
"""
import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_conga_template(output_path):
    """
    Create a sample Conga template with merge fields for testing.
    
    Args:
        output_path: Path to save the sample template
    """
    # Create a new document
    doc = docx.Document()
    
    # Add a title
    title = doc.add_heading('Sample Contract', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add company information section
    doc.add_heading('Company Information', 1)
    p = doc.add_paragraph('Company Name: ')
    p.add_run('«Company_Name»').bold = True
    
    p = doc.add_paragraph('Address: ')
    p.add_run('«Company_Street»').bold = True
    p.add_run(', ')
    p.add_run('«Company_City»').bold = True
    p.add_run(', ')
    p.add_run('«Company_State»').bold = True
    p.add_run(' ')
    p.add_run('«Company_Zip»').bold = True
    
    # Add customer information section
    doc.add_heading('Customer Information', 1)
    p = doc.add_paragraph('Customer: ')
    p.add_run('«Contact_FirstName»').bold = True
    p.add_run(' ')
    p.add_run('«Contact_LastName»').bold = True
    
    p = doc.add_paragraph('Email: ')
    p.add_run('«Contact_Email»').bold = True
    
    p = doc.add_paragraph('Phone: ')
    p.add_run('«Contact_Phone»').bold = True
    
    # Add contract details section
    doc.add_heading('Contract Details', 1)
    p = doc.add_paragraph('Contract Number: ')
    p.add_run('«Contract_Number»').bold = True
    
    p = doc.add_paragraph('Start Date: ')
    p.add_run('«Contract_StartDate»').bold = True
    
    p = doc.add_paragraph('End Date: ')
    p.add_run('«Contract_EndDate»').bold = True
    
    p = doc.add_paragraph('Amount: $')
    p.add_run('«Contract_Amount»').bold = True
    
    # Add products section
    doc.add_heading('Products', 1)
    
    # Create a table for products
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Product Name'
    header_cells[1].text = 'Quantity'
    header_cells[2].text = 'Unit Price'
    header_cells[3].text = 'Total'
    
    # Add a sample product row with merge fields
    row_cells = table.add_row().cells
    row_cells[0].text = '«Product_Name»'
    row_cells[1].text = '«Product_Quantity»'
    row_cells[2].text = '$«Product_UnitPrice»'
    row_cells[3].text = '$«Product_Total»'
    
    # Add signature section
    doc.add_heading('Signatures', 1)
    p = doc.add_paragraph('Customer Signature: ________________________')
    p = doc.add_paragraph('Date: ________________')
    p = doc.add_paragraph()
    p = doc.add_paragraph('Company Representative: ')
    p.add_run('«Rep_Name»').bold = True
    p = doc.add_paragraph('Date: ________________')
    
    # Save the document
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    # Create sample data directory if it doesn't exist
    sample_dir = "/home/ubuntu/conga_to_box_converter/sample_data"
    os.makedirs(sample_dir, exist_ok=True)
    
    # Create sample template
    template_path = os.path.join(sample_dir, "sample_template.docx")
    create_sample_conga_template(template_path)
    
    print(f"Sample Conga template created at: {template_path}")
